import os, os.path, errno, sys, traceback, subprocess
import re, html.entities
import json
import logging
import yaml
from bs4 import BeautifulSoup
from datetime import datetime
import requests
import urllib.parse
import io
import gzip
import certifi
import docx
import zipfile
from urllib.parse import urljoin
import inspect
import pdfrw

from . import admin

logging.getLogger("pdfrw").setLevel(logging.CRITICAL)

# scraper should be instantiated at class-load time, so that it can rate limit appropriately
import scrapelib
scraper = scrapelib.Scraper(requests_per_minute=120, retry_attempts=3)
scraper.user_agent = "unitedstates/inspectors-general (https://github.com/unitedstates/inspectors-general)"
scraper.timeout = 60

class Soft404HttpAdapter(requests.adapters.HTTPAdapter):
  """Transport adapter that checks all responses against a blacklist of "file
  not found" pages that are served with 200 status codes."""

  SOFT_404_URLS_RE = re.compile(r"^(http://www\.dodig\.mil/errorpages/index\.html|http://www\.fec\.gov/404error\.shtml|http://www\.gpo\.gov/maintenance/error\.htm)$")
  SOFT_404_BODY_SIGNATURES = {
    "cftc.gov": b"<title>404 Page Not Found - CFTC</title>",
    "cpb.org": b"<title>CPB: Page Not Found</title>",
    "ncua.gov": b"Redirect.aspx?404",
    "si.edu": b"<title>Page Not Found Smithsonian</title>",
  }

  def build_response(self, req, resp):
    domain = urllib.parse.urlparse(req.url)[1].split(':')[0]
    base_domain = ".".join(domain.split(".")[-2:])
    if base_domain in self.SOFT_404_BODY_SIGNATURES:
      if resp.getheader("Content-Type") in ["text/html; charset=utf-8",
                                            "text/html"]:
        data = resp.data
        headers = resp.headers
        if resp.getheader("Content-Encoding") == "gzip":
          decompressed_data = gzip.decompress(data)
        else:
          decompressed_data = data
        if resp.getheader("Transfer-Encoding") == "chunked":
          headers.pop("Transfer-Encoding")
        body = io.BytesIO(data)
        resp = requests.packages.urllib3.response.HTTPResponse(
                body=body,
                headers=headers,
                status=resp.status,
                version=resp.version,
                reason=resp.reason,
                strict=resp.strict,
                preload_content=False,
        )
        if decompressed_data.find(self.SOFT_404_BODY_SIGNATURES[base_domain], 0, 10240) != -1:
          result = super(Soft404HttpAdapter, self).build_response(req, resp)
          result.status_code = 404 # tells scrapelib to not retry
          return result

    redirect = resp.get_redirect_location()
    result = super(Soft404HttpAdapter, self).build_response(req, resp)
    if redirect and self.SOFT_404_URLS_RE.match(redirect):
      result.status_code = 404 # tells scrapelib to not retry

    return result

scraper.mount("http://www.cftc.gov/", Soft404HttpAdapter())
scraper.mount("http://cftc.gov/", Soft404HttpAdapter())
scraper.mount("http://www.dodig.mil/", Soft404HttpAdapter())
scraper.mount("http://dodig.mil/", Soft404HttpAdapter())
scraper.mount("http://www.fec.gov/", Soft404HttpAdapter())
scraper.mount("http://fec.gov/", Soft404HttpAdapter())
scraper.mount("http://www.gpo.gov/", Soft404HttpAdapter())
scraper.mount("http://gpo.gov/", Soft404HttpAdapter())
scraper.mount("http://www.cpb.org/", Soft404HttpAdapter())
scraper.mount("http://cpb.org/", Soft404HttpAdapter())
scraper.mount("http://www.ncua.gov/", Soft404HttpAdapter())
scraper.mount("http://ncua.gov/", Soft404HttpAdapter())
scraper.mount("http://www.si.edu/", Soft404HttpAdapter())
scraper.mount("http://si.edu/", Soft404HttpAdapter())

class CipherListAdapter(requests.adapters.HTTPAdapter):
  def __init__(self, ciphers):
    self.ciphers = ciphers
    super(CipherListAdapter, self).__init__()

  def init_poolmanager(self, num_pools, maxsize, block=False, *args, **kwargs):
    context = requests.packages.urllib3.util.ssl_.create_urllib3_context(
        ciphers=self.ciphers
    )
    kwargs["ssl_context"] = context
    self.poolmanager = requests.packages.urllib3.poolmanager.PoolManager(
        num_pools=num_pools,
        maxsize=maxsize,
        block=block,
        *args,
        **kwargs
    )

# The ARC server or middlebox only supports one cipher suite,
# DES-CBC3-SHA, and it also needs it to be sufficiently far forward in the
# cipher suite list.
# (as of 9/8/2016)
scraper.mount("https://www.arc.gov/", CipherListAdapter("DES-CBC3-SHA"))

WHITELIST_INSECURE_DOMAINS = (
  "https://www.ignet.gov/",  # incomplete chain as of 1/25/2015
  "https://www.va.gov/",  # incomplete chain as of 12/6/2015
  "https://origin.www.fhfaoig.gov/",  # incomplete chain as of 1/5/2016
  "https://www.archives.gov/",  # incomplete chain as of 2/19/2016
  "https://www.arc.gov/",  # incomplete chain as of 8/24/2016
  "https://www.ncua.gov/",  # incomplete chain as of 12/10/2016
  "https://www.oig.dot.gov/",  # incomplete chain as of 4/12/2017
  "https://www.oig.dhs.gov",  # incomplete chain as of 5/6/2017
  "https://oig.eeoc.gov/",  # incomplete chain as of 6/3/2017
  "https://www.cpb.org/",  # incomplete chain as of 10/1/2017

  # The following domains will 301/302 redirect to the above domains, so
  # validate=False is needed for these cases as well
  "http://www.ignet.gov/",
  "http://ignet.gov/",
  "http://www.arc.gov/",
  "http://www.cpb.org/",
)
WHITELIST_SHA1_DOMAINS = (
  "https://www.sba.gov/",
  "http://www.sba.gov/",
  "http://sba.gov/",
  "https://www.sigar.mil/",
  "http://www.sigar.mil/",
  "https://oig.nasa.gov/",
  "https://www.gsaig.gov/",
  "https://www.house.gov/",
)

# Special case handling for governmentattic.org, va.gov, etc.:
# These pages are served without an encoding in the HTTP headers,
# and with the encoding specified in a <meta> tag inside the document.
# See https://gist.github.com/divergentdave/0985ae1f2fd2a7235ccef0d5bbb6aaa3
# for data collection script.
META_CHARSETS = {
  "http://www.usda.gov/oig": "iso-8859-1",
  "http://www.cftc.gov/": "utf-8",
  "http://cftc.gov/": "utf-8",
  "https://www.cftc.gov/": "utf-8",
  "https://cftc.gov/": "utf-8",
  "http://www.dodig.mil/": "utf-8",
  "https://www.dodig.mil/": "utf-8",
  "https://oig.justice.gov/": "utf-8",
  "https://www.fca.gov/": "utf-8",
  "https://fca.gov/": "utf-8",
  "http://www.fec.gov/": "iso-8859-1",
  "https://www.fec.gov/": "iso-8859-1",
  "https://oig.federalreserve.gov/": "iso-8859-1",
  "http://www.gao.gov/": "utf-8",
  "http://gao.gov/": "utf-8",
  "https://www.gao.gov/": "utf-8",
  "https://gao.gov/": "utf-8",
  "http://www.governmentattic.org/": "utf-8",
  "http://governmentattic.org/": "utf-8",
  "https://www.governmentattic.org/": "utf-8",
  "https://governmentattic.org/": "utf-8",
  "https://oig.nasa.gov/": "utf-8",
  "http://www.nrc.gov/": "utf-8",
  "https://www.nrc.gov/": "utf-8",
  "http://oig.pbgc.gov/": "utf-8",
  "https://oig.pbgc.gov/": "utf-8",
  "https://www.treasury.gov/tigta/": "iso-8859-1",
  "http://oig.tva.gov/": "utf-8",
  "https://oig.tva.gov/": "utf-8",
  "https://www.va.gov/oig/": "utf-8",
}

# will pass correct options on to individual scrapers whether
# run through ./igs or individually, because argv[1:] is the same
def run(run_method, additional=None):
  cli_options = options()
  configure_logging(cli_options)

  if additional:
    cli_options.update(additional)

  try:
    return run_method(cli_options)
  except Exception as exception:
    admin.log_exception(exception)


# read options from the command line
#   e.g. ./inspectors/usps.py --since=2012-03-04 --debug
#     => {"since": "2012-03-04", "debug": True}
AVAILABLE_OPTIONS = (
  "archive",
  "bulk",
  "component",
  "debug",
  "dry_run",
  "end",
  "ig",
  "limit",
  "log",
  "only",
  "pages",
  "quick",
  "report_id",
  "safe",
  "since",
  "skip_downloaded",
  "start",
  "topics",
  "types",
  "year",
)


def options():
  options = {}
  for arg in sys.argv[1:]:
    if arg.startswith("--"):

      if "=" in arg:
        key, value = arg.split('=')
      else:
        key, value = arg, "true"

      key = key.split("--")[1]
      key = key.lower()
      value = value.lower()

      if key not in AVAILABLE_OPTIONS:
        print("Unknown option: \"%s\"\n"
              "The following options are recognized\n"
              "  %s"% (key, ", ".join(AVAILABLE_OPTIONS)))
        sys.exit(1)

      if value == 'true': value = True
      elif value == 'false': value = False
      options[key] = value

  return options

def configure_logging(options=None):
  options = {} if not options else options
  if options.get('debug', False):
    log_level = "debug"
  else:
    log_level = options.get("log", "warn")

  if log_level not in ["debug", "info", "warn", "error"]:
    print("Invalid log level (specify: debug, info, warn, error).")
    sys.exit(1)

  logging.basicConfig(format='%(message)s', level=log_level.upper())


# download the data at url
def download(url, destination=None, options=None, scraper_slug=None):
  options = {} if not options else options
  cache = options.get('cache', True) # default to caching
  binary = options.get('binary', False) # default to assuming text

  # check cache first
  if destination and cache and os.path.exists(destination):
    logging.info("## Cached: (%s, %s)" % (destination, url))

    # if a binary file is cached, we're done
    if binary:
      return True

    # otherwise, decode it for return
    with open(destination, 'r', encoding='utf-8') as f:
      body = f.read()

  # otherwise, download from the web
  else:

    logging.warn(url)

    logging.info("## Downloading: %s" % url)
    if binary:
      if destination:
        logging.info("## \tto: %s" % destination)
      else:
        raise Exception("A destination path is required for downloading a binary file")
      try:
        mkdir_p(os.path.dirname(destination))

        verify_options = domain_verify_options(url)
        scraper.urlretrieve(url, destination, verify=verify_options)
      except connection_errors() as e:
        admin.log_http_error(e, url, scraper_slug)
        return None
    else: # text
      try:
        if destination: logging.info("## \tto: %s" % destination)

        # Special handler for downloading reports whose server has
        # misconfigured their HTTPS, and for which no alternative
        # exists.
        # This happens very rarely, and scrapelib has a bug with
        # verification options, so this disables the rate limiting
        # provided by scrapelib.

        verify_options = domain_verify_options(url)
        response = scraper.get(url, verify=verify_options)

      except connection_errors() as e:
        admin.log_http_error(e, url, scraper_slug)
        return None

      for prefix, charset in META_CHARSETS.items():
        if url.startswith(prefix):
          response.encoding = charset
          break

      body = response.text
      if not isinstance(body, str): raise ValueError("Content not decoded.")

      # don't allow 0-byte files
      if (not body) or (not body.strip()):
        return None

      # cache content to disk
      if destination:
        write(body, destination, binary=binary)

  # don't return binary content
  if binary:
    return True
  else:
    # whether from disk or web, unescape HTML entities
    return unescape(body)

def beautifulsoup_from_url(url):
  caller_filename = inspect.stack()[1][1]
  caller_scraper = os.path.splitext(os.path.basename(caller_filename))[0]

  body = download(url, scraper_slug=caller_scraper)
  if body is None: return None

  doc = BeautifulSoup(body, "lxml")

  # Some of the pages will return meta refreshes
  if doc.find("meta") and doc.find("meta").attrs.get('http-equiv') == 'REFRESH':
    redirect_url = urljoin(url, doc.find("meta").attrs['content'].split("url=")[1])
    return beautifulsoup_from_url(redirect_url)
  else:
    return doc

def post(url, data=None, headers=None, **kwargs):
  response = None
  try:
    verify_options = domain_verify_options(url)
    response = scraper.post(url, data=data, headers=headers, verify=verify_options)
  except connection_errors() as e:
    admin.log_http_error(e, url)
    return None

  return response

def resolve_redirect(url):
  res = scraper.request(method='HEAD', url=url, allow_redirects=False)
  if "Location" in res.headers:
    return res.headers["Location"]
  else:
    return url

def connection_errors():
  return (scrapelib.HTTPError, requests.exceptions.ConnectionError, requests.packages.urllib3.exceptions.MaxRetryError)

# uses BeautifulSoup to do a naive extraction of text from HTML,
# then writes it and returns the /data-relative path.
def text_from_html(real_html_path, real_text_path):
  html = open(real_html_path, encoding='utf-8').read()
  doc = BeautifulSoup(html, "lxml")

  for node in doc.findAll(['script', 'style']):
    node.extract()

  text = doc.text
  lines = text.splitlines()
  for i in range(len(lines)):
    lines[i] = lines[i].strip()
  lines = filter(None, lines)
  text = "\n".join(lines)

  write(text, real_text_path, binary=False)

def domain_verify_options(url):
  for domain in WHITELIST_SHA1_DOMAINS:
    if url.startswith(domain):
      return certifi.old_where()
  for domain in WHITELIST_INSECURE_DOMAINS:
    if url.startswith(domain):
      logging.warn("SKIPPING HTTPS VERIFICATION.")
      return False
  return True


_tool_present_cache = {}


def check_tool_present(*args):
  if args in _tool_present_cache:
    return _tool_present_cache[args]
  try:
    subprocess.Popen(args,
                     shell=False,
                     stdout=subprocess.DEVNULL,
                     stderr=subprocess.STDOUT).communicate()
    result = True
  except FileNotFoundError:
    result = False
  _tool_present_cache[args] = result
  return result

# read PDF's directory to determine if we need to decrypt it
def check_pdf_decryption(pdf_path):
  try:
    doc = pdfrw.PdfReader(pdf_path)
    return "/Encrypt" in doc
  except:
    return False

# uses qpdf to decrypt a PDF
def decrypt_pdf(source_path, destination_path):
  if not check_tool_present("qpdf", "--version"):
    logging.warn("Install qpdf to decrypt PDFs! "
                 "The qpdf executable must be in a directory that is in "
                 "your PATH environment variable.")
    return False

  try:
    subprocess.check_call(["qpdf",
                           "--decrypt",
                           source_path,
                           destination_path], shell=False)
    return True
  except subprocess.CalledProcessError as exc:
    logging.warn("Error decrypting %s:\n\n%s" %
                 (source_path, format_exception(exc)))
    return False

  if not os.path.exists(destination_path):
    logging.warn("PDF not decrypted to %s" % destination_path)
    return False

# uses pdftotext to get text out of PDFs,
# then writes it and returns the /data-relative path.
def text_from_pdf(real_pdf_path, real_text_path):
  if not check_tool_present("pdftotext", "-v"):
    logging.warn("Install pdftotext to extract text! "
                 "The pdftotext executable must be in a directory that is in "
                 "your PATH environment variable.")
    return

  try:
    subprocess.check_call(["pdftotext",
                           "-layout",
                           "-nopgbrk",
                           real_pdf_path,
                           real_text_path], shell=False)
  except subprocess.CalledProcessError as exc:
    logging.warn("Error extracting text to %s:\n\n%s" %
                 (real_text_path, format_exception(exc)))
    return

  if not os.path.exists(real_text_path):
    logging.warn("Text not extracted to %s" % real_text_path)

def text_from_doc(real_doc_path, real_text_path):
  if not check_tool_present("abiword", "-?"):
    logging.warn("Install AbiWord to extract text! "
                 "The abiword executable must be in a directory that is in "
                 "your PATH environment variable.")
    return

  try:
    subprocess.check_call(["abiword",
                           real_doc_path,
                           "--to",
                           "txt"], shell=False)
  except subprocess.CalledProcessError as exc:
    logging.warn("Error extracting text to %s:\n\n%s" %
                 (real_text_path, format_exception(exc)))
    return

  if not os.path.exists(real_text_path):
    logging.warn("Text not extracted to %s" % real_text_path)

def text_from_docx(real_docx_path, real_text_path):
  def text_from_paragraphs(paragraphs):
    return "\n\n".join([paragraph.text for paragraph in paragraphs])

  def text_from_tables(tables):
    return "\n\n".join([text_from_table(table) for table in tables])

  def text_from_table(table):
    return "\n\n".join([text_from_row(row) for row in table.rows])

  def text_from_row(row):
    return "\n\n".join([text_from_doc_or_cell(cell) for cell in row.cells])

  def text_from_doc_or_cell(cell):
    part1 = text_from_paragraphs(cell.paragraphs)
    part2 = text_from_tables(cell.tables)
    if part1 and part2:
      return "%s\n\n%s" % (part1, part2)
    else:
      return part1 + part2

  try:
    document = docx.Document(real_docx_path)
    text = text_from_doc_or_cell(document)
    write(text, real_text_path, binary=False)
  except zipfile.BadZipFile as exc:
    logging.warn("Error extracting text to %s:\n\n%s" %
                 (real_text_path, format_exception(exc)))
    return None

PDF_PAGE_RE = re.compile("Pages: +([0-9]+)\r?\n")
PDF_CREATION_DATE_RE = re.compile("CreationDate: +([^\r\n]*)\r?\n")
PDF_MOD_DATE_RE = re.compile("ModDate: +([^\r\n]*)\r?\n")
PDF_TITLE_RE = re.compile("Title: +([^\r\n]*)\r?\n")
PDF_KEYWORDS_RE = re.compile("Keywords: +([^\r\n]*)\r?\n")
PDF_AUTHOR_RE = re.compile("Author: +([^\r\n]*)\r?\n")

def parse_pdf_datetime(raw):
  if raw.strip() == "":
    return None
  my_datetime = None

  datetime_formats = [
    '%m/%d/%y %H:%M:%S',
    '%a %b %d %H:%M:%S %Y',
    '%A, %B %d, %Y %I:%M:%S %p',
    '%a %b %d %H:%M:%S %Y %Z'
  ]
  for datetime_format in datetime_formats:
    try:
      my_datetime = datetime.strptime(raw, datetime_format)
      break
    except ValueError:
      pass

  if my_datetime:
    return datetime.strftime(my_datetime, '%Y-%m-%d')
  else:
    logging.warn('Could not parse PDF date: %s' % raw)
    return None

def metadata_from_pdf(pdf_path):
  if not check_tool_present("pdfinfo", "-v"):
    logging.warn("Install pdfinfo to extract metadata! "
                 "The pdfinfo executable must be in a directory that is in "
                 "your PATH environment variable.")
    return None

  real_pdf_path = os.path.expandvars(os.path.join(data_dir(), pdf_path))
  real_pdf_path = os.path.abspath(real_pdf_path)

  try:
    output = subprocess.check_output(["pdfinfo", real_pdf_path], shell=False)
    output = output.decode('utf-8', errors='replace')
  except subprocess.CalledProcessError as exc:
    logging.warn("Error extracting metadata for %s:\n\n%s" %
                 (pdf_path, format_exception(exc)))
    return None

  metadata = {}

  page_match = PDF_PAGE_RE.search(output)
  if page_match:
    metadata['page_count'] = int(page_match.group(1))

  creation_date_match = PDF_CREATION_DATE_RE.search(output)
  if creation_date_match:
    metadata['creation_date'] = parse_pdf_datetime(creation_date_match.group(1))

  mod_date_match = PDF_MOD_DATE_RE.search(output)
  if mod_date_match:
    metadata['modification_date'] = parse_pdf_datetime(mod_date_match.group(1))

  title_match = PDF_TITLE_RE.search(output)
  if title_match:
    metadata['title'] = title_match.group(1)

  keywords_match = PDF_KEYWORDS_RE.search(output)
  if keywords_match:
    metadata['keywords'] = keywords_match.group(1)

  author_match = PDF_AUTHOR_RE.search(output)
  if author_match:
    metadata['author'] = author_match.group(1)

  if metadata:
    return metadata
  return None

def check_report_url(report_url):
  try:
    verify_options = domain_verify_options(report_url)
    scraper.request(method='HEAD', url=report_url, verify=verify_options)
  except connection_errors() as e:
    admin.log_http_error(e, report_url)

DOC_PAGE_RE = re.compile("Number of Pages: ([0-9]*),")
DOC_CREATION_DATE_RE = re.compile("Create Time/Date: ([A-Za-z 0-9:]*),")
DOC_MOD_DATE_RE = re.compile("Last Saved Time/Date: ([A-Za-z 0-9:]*),")
DOC_TITLE_RE = re.compile("Title: ([^,]*),")
DOC_AUTHOR_RE = re.compile("Author: ([^,]*),")

def parse_doc_datetime(raw):
  if raw.strip() == "":
    return None
  my_datetime = None
  try:
    my_datetime = datetime.strptime(raw, '%a %b %d %H:%M:%S %Y')
  except ValueError:
    pass
  if my_datetime:
    return datetime.strftime(my_datetime, '%Y-%m-%d')
  else:
    logging.warn('Could not parse DOC date: %s' % raw)
    return None

def metadata_from_doc(doc_path):
  if not check_tool_present("file", "-v"):
    logging.warn("Install file to extract metadata! "
                 "The file executable must be in a directory that is in your "
                 "PATH environment variable.")
    return None

  real_doc_path = os.path.expandvars(os.path.join(data_dir(), doc_path))
  real_doc_path = os.path.abspath(real_doc_path)

  try:
    output = subprocess.check_output(["file", real_doc_path], shell=False)
    output = output.decode('utf-8', errors='replace')
  except subprocess.CalledProcessError as exc:
    logging.warn("Error extracting metadata for %s:\n\n%s" %
                 (doc_path, format_exception(exc)))
    return None

  metadata = {}

  page_match = DOC_PAGE_RE.search(output)
  if page_match:
    metadata['page_count'] = int(page_match.group(1))

  creation_date_match = DOC_CREATION_DATE_RE.search(output)
  if creation_date_match:
    metadata['creation_date'] = parse_doc_datetime(creation_date_match.group(1))

  mod_date_match = DOC_MOD_DATE_RE.search(output)
  if mod_date_match:
    metadata['mod_date'] = parse_doc_datetime(mod_date_match.group(1))

  title_match = DOC_TITLE_RE.search(output)
  if title_match:
    metadata['title'] = title_match.group(1)

  author_match = DOC_AUTHOR_RE.search(output)
  if author_match:
    metadata['author'] = author_match.group(1)

  if metadata:
    return metadata
  return None

def metadata_from_docx(docx_path):
  try:
    real_docx_path = os.path.expandvars(os.path.join(data_dir(), docx_path))
    real_docx_path = os.path.abspath(real_docx_path)
    document = docx.Document(real_docx_path)
    core_props = document.core_properties

    metadata = {}

    if core_props.author:
      metadata['author'] = core_props.author

    if core_props.title:
      metadata['title'] = core_props.title

    if core_props.created:
      metadata['creation_date'] = datetime.strftime(core_props.created, '%Y-%m-%d')

    if core_props.modified:
      metadata['mod_date'] = datetime.strftime(core_props.created, '%Y-%m-%d')

    if core_props.keywords:
      metadata['keywords'] = core_props.keywords

    if metadata:
      return metadata
    return None
  except zipfile.BadZipFile as exc:
    logging.warn("Error extracting metadata for %s:\n\n%s" %
                 (docx_path, format_exception(exc)))
    return None

def format_exception(exception):
  exc_type, exc_value, exc_traceback = sys.exc_info()
  return "\n".join(traceback.format_exception(exc_type, exc_value, exc_traceback))

# assumes working dir is the root dir
def data_dir():
  if admin.config and admin.config.get('data_directory'):
    return admin.config.get('data_directory')
  return "data"

def write(content, destination, binary=False):
  mkdir_p(os.path.dirname(destination))

  if binary:
    f = open(destination, 'bw')
  else:
    f = open(destination, 'w', encoding='utf-8')
  f.write(content)
  f.close()

def json_for(object):
  return json.dumps(object, sort_keys=True, indent=2, default=format_datetime)

def format_datetime(obj):
  if isinstance(obj, datetime.date):
    return obj.isoformat()
  elif isinstance(obj, str):
    return obj
  else:
    return None

# mkdir -p in python, from:
# http://stackoverflow.com/questions/600268/mkdir-p-functionality-in-python
def mkdir_p(path):
  try:
    os.makedirs(path)
  except OSError as exc: # Python >2.5
    if exc.errno == errno.EEXIST:
      pass
    else:
      raise

# taken from http://effbot.org/zone/re-sub.htm#unescape-html
def unescape(text):

  def remove_unicode_control(str):
    remove_re = re.compile('[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]')
    return remove_re.sub('', str)

  def fixup(m):
    text = m.group(0)
    if text[:2] == "&#":
      # character reference
      try:
        if text[:3] == "&#x":
          return chr(int(text[3:-1], 16))
        else:
          return chr(int(text[2:-1]))
      except ValueError:
        pass
    else:
      # named entity
      try:
        text = chr(html.entities.name2codepoint[text[1:-1]])
      except KeyError:
        pass
    return text # leave as is

  text = re.sub("&#?\w+;", fixup, text)
  text = remove_unicode_control(text)
  return text

# 'safe' scrapers listed in safe.yml
def safe_igs():
  return yaml.load(open("safe.yml"))
